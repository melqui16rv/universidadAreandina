"""
Generador DOCX para Eje 3.

Convierte `documento_entrega.md` a Word y busca imagenes en `img`.
Replica el flujo de Eje 1 en formato individual para Registro + FTK Imager.
"""

from pathlib import Path
import re

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


BASE_DIR = Path(__file__).resolve().parent
MD_FILE = BASE_DIR / "documento_entrega.md"
IMG_DIR = BASE_DIR / "img"
LEGACY_IMG_DIR = BASE_DIR.parent / "img"
OUTPUT_FILE = BASE_DIR / "Analisis_Forense_Registro_Windows_MelquiRomero.docx"


def normalize_line(line: str) -> str:
	return line.rstrip("\n").rstrip("\r")


def is_table_line(line: str) -> bool:
	return line.strip().startswith("|") and line.strip().endswith("|")


def split_table_row(line: str):
	parts = [p.strip() for p in line.strip().strip("|").split("|")]
	return parts


def set_doc_style(doc: Document):
	section = doc.sections[0]
	section.left_margin = Cm(2.54)
	section.right_margin = Cm(2.54)
	section.top_margin = Cm(2.54)
	section.bottom_margin = Cm(2.54)

	style = doc.styles["Normal"]
	style.font.name = "Calibri"
	style.font.size = Pt(11)


def add_missing_image_marker(doc: Document, img_name: str):
	p = doc.add_paragraph()
	run = p.add_run(f"[Imagen no encontrada: {img_name}]")
	run.italic = True
	run.font.color.rgb = RGBColor(180, 0, 0)
	p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def insert_image_from_markdown(doc: Document, line: str):
	# Soporta: ![texto](ruta)
	match = re.match(r"!\[[^\]]*\]\(([^\)]+)\)", line.strip())
	if not match:
		return False

	raw_path = match.group(1).strip()
	candidate = (BASE_DIR / raw_path).resolve()
	if not candidate.exists():
		candidate = (IMG_DIR / Path(raw_path).name).resolve()
	if not candidate.exists():
		candidate = (LEGACY_IMG_DIR / Path(raw_path).name).resolve()

	if candidate.exists():
		doc.add_picture(str(candidate), width=Inches(6.0))
		doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
	else:
		add_missing_image_marker(doc, Path(raw_path).name)

	return True


def parse_inline_bold(paragraph, text: str):
	# Manejo simple de **bold**
	chunks = re.split(r"(\*\*[^*]+\*\*)", text)
	for chunk in chunks:
		if chunk.startswith("**") and chunk.endswith("**") and len(chunk) > 4:
			run = paragraph.add_run(chunk[2:-2])
			run.bold = True
		else:
			paragraph.add_run(chunk)


def add_table(doc: Document, rows):
	if not rows:
		return

	# Detecta y remueve linea separadora markdown |---|
	parsed_rows = [split_table_row(r) for r in rows]
	if len(parsed_rows) >= 2 and all(set(col.replace("-", "").replace(":", "").strip()) == set() for col in parsed_rows[1]):
		parsed_rows.pop(1)

	col_count = max(len(r) for r in parsed_rows)
	table = doc.add_table(rows=len(parsed_rows), cols=col_count)
	table.style = "Table Grid"
	table.alignment = WD_TABLE_ALIGNMENT.CENTER

	for i, row in enumerate(parsed_rows):
		for j in range(col_count):
			txt = row[j] if j < len(row) else ""
			table.cell(i, j).text = ""
			cell_p = table.cell(i, j).paragraphs[0]
			run = cell_p.add_run(txt)
			run.font.size = Pt(10)
			if i == 0:
				run.bold = True


def build_doc_from_markdown(md_text: str) -> Document:
	doc = Document()
	set_doc_style(doc)

	lines = [normalize_line(l) for l in md_text.splitlines()]
	i = 0
	while i < len(lines):
		line = lines[i]
		stripped = line.strip()

		# Saltos
		if stripped == "":
			doc.add_paragraph("")
			i += 1
			continue

		# Imagenes
		if insert_image_from_markdown(doc, stripped):
			i += 1
			continue

		# Titulos
		if stripped.startswith("### "):
			h = doc.add_heading(stripped[4:].strip(), level=3)
			h.runs[0].font.name = "Calibri"
			i += 1
			continue
		if stripped.startswith("## "):
			h = doc.add_heading(stripped[3:].strip(), level=2)
			h.runs[0].font.name = "Calibri"
			i += 1
			continue
		if stripped.startswith("# "):
			h = doc.add_heading(stripped[2:].strip(), level=1)
			h.runs[0].font.name = "Calibri"
			i += 1
			continue

		# Separador markdown ---
		if stripped == "---":
			p = doc.add_paragraph("_" * 70)
			p.alignment = WD_ALIGN_PARAGRAPH.CENTER
			i += 1
			continue

		# Tablas
		if is_table_line(stripped):
			table_lines = []
			while i < len(lines) and is_table_line(lines[i]):
				table_lines.append(lines[i])
				i += 1
			add_table(doc, table_lines)
			continue

		# Listas con guion
		if stripped.startswith("- "):
			p = doc.add_paragraph(style="List Bullet")
			parse_inline_bold(p, stripped[2:].strip())
			i += 1
			continue

		# Listas numeradas
		num_match = re.match(r"^\d+\.\s+(.*)$", stripped)
		if num_match:
			p = doc.add_paragraph(style="List Number")
			parse_inline_bold(p, num_match.group(1).strip())
			i += 1
			continue

		# Parrafo normal
		p = doc.add_paragraph()
		parse_inline_bold(p, stripped)
		i += 1

	return doc


def main():
	if not MD_FILE.exists():
		raise FileNotFoundError(f"No existe el archivo markdown: {MD_FILE}")

	md_text = MD_FILE.read_text(encoding="utf-8")
	doc = build_doc_from_markdown(md_text)
	doc.save(str(OUTPUT_FILE))
	print(f"DOCX generado correctamente: {OUTPUT_FILE}")


if __name__ == "__main__":
	main()

