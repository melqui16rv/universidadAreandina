"""
Script para generar el documento Word (.docx) a partir del contenido Markdown
del análisis forense de eventos de Windows.
"""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ─── Configuración ───────────────────────────────────────────────────────────

BASE_DIR = Path(__file__).parent
IMG_DIR = BASE_DIR / "img"
MD_FILE = BASE_DIR / "documento_entrega.md"
OUTPUT_FILE = BASE_DIR / "Analisis_Forense_Eventos_Windows_MelquiRomero.docx"

# ─── Utilidades ──────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda de tabla."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)


def add_formatted_run(paragraph, text, bold=False, italic=False, size=None, color=None, font_name=None):
    """Agrega un run con formato al párrafo."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    return run


def add_image_safe(doc, img_path, width=Inches(5.8)):
    """Inserta una imagen si existe, o un texto indicando que falta."""
    if img_path.exists():
        p = doc.add_picture(str(img_path), width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        add_formatted_run(p, f"[Imagen no encontrada: {img_path.name}]", italic=True, color=RGBColor(200, 0, 0))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc, headers, rows, header_color="1F4E79", col_widths=None):
    """Crea una tabla formateada."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezados
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        add_formatted_run(p, header, bold=True, size=9, color=RGBColor(255, 255, 255), font_name="Calibri")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # Filas de datos
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[1 + r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            # Si texto contiene ⚠️ o 🔴, aplicar color
            if "⚠️" in cell_text:
                add_formatted_run(p, cell_text, size=8, color=RGBColor(200, 100, 0), font_name="Calibri")
            elif "🔴" in cell_text or "Crítico" in cell_text or "Alto" in cell_text:
                add_formatted_run(p, cell_text, size=8, color=RGBColor(200, 0, 0), font_name="Calibri")
            else:
                add_formatted_run(p, cell_text, size=8, font_name="Calibri")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Fondo alterno para legibilidad
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F2F2")

    return table


def add_key_value_table(doc, pairs, header_color="2E75B6"):
    """Crea una tabla de campo-valor (2 columnas)."""
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (key, val) in enumerate(pairs):
        cell_k = table.rows[i].cells[0]
        cell_v = table.rows[i].cells[1]
        cell_k.text = ""
        cell_v.text = ""
        p_k = cell_k.paragraphs[0]
        p_v = cell_v.paragraphs[0]
        add_formatted_run(p_k, key, bold=True, size=9, font_name="Calibri")
        add_formatted_run(p_v, val, size=9, font_name="Calibri")
        set_cell_shading(cell_k, "D6E4F0")

    return table


# ─── Construcción del Documento ──────────────────────────────────────────────

def build_document():
    doc = Document()

    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # ═══════════════════════════════════════════════════════════════════════
    # PORTADA
    # ═══════════════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(title, "Análisis Forense de un Registro\nde Eventos en Windows", bold=True, size=26, color=RGBColor(31, 78, 121))

    doc.add_paragraph("")

    portada_data = [
        ("Asignatura:", "Informática Forense"),
        ("Eje:", "Eje 1"),
        ("Estudiante:", "Melqui Romero"),
        ("Docente:", "(por definir)"),
        ("Institución:", "SENA"),
        ("Fecha:", "Febrero de 2026"),
    ]
    add_key_value_table(doc, portada_data)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 1. INTRODUCCIÓN
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("1. Introducción", level=1)

    intro_paragraphs = [
        'La informática forense es la disciplina que aplica técnicas científicas y analíticas especializadas para identificar, preservar, analizar y presentar datos que sean válidos dentro de un proceso legal o investigativo (Luna Felipez, s.f.). En un mundo cada vez más digitalizado, donde los sistemas informáticos son el eje de las comunicaciones, transacciones y operaciones cotidianas, el análisis forense digital se ha convertido en una herramienta imprescindible para la investigación de incidentes de ciberseguridad y cibercrimen.',
        'El principio de intercambio de Locard, piedra angular de la ciencia forense, establece que "todo contacto deja un rastro". En el contexto digital, esto significa que cada vez que un usuario interactúa con un sistema informático —ya sea un computador, una red o un dispositivo móvil— se generan rastros que evidencian dicha interacción: archivos, logs de actividad, metadatos, historiales de navegación y correos electrónicos, entre otros (De Haro Olmo, 2020). Estos rastros digitales constituyen la evidencia que un perito informático forense debe identificar, preservar y analizar con rigor metodológico.',
        'El sistema operativo Windows, al ser uno de los más utilizados a nivel mundial, integra una herramienta nativa denominada Visor de Eventos (Event Viewer) que registra de manera automática y continua toda la actividad del sistema: inicios de sesión exitosos y fallidos, cambios de privilegios, instalación de servicios, errores del sistema y mucho más. Estos registros (logs) son una fuente primaria de evidencia digital en cualquier investigación forense.',
        'El presente trabajo tiene como propósito realizar un análisis forense de los registros de eventos de un sistema Windows, simulando un escenario de investigación donde se busca identificar posibles huellas de actividad sospechosa o maliciosa. Se aplicarán los conceptos del principio de intercambio de Locard y las fases de la metodología forense digital: identificación, preservación, análisis y presentación de la evidencia.',
    ]
    for text in intro_paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(1.27)

    # ═══════════════════════════════════════════════════════════════════════
    # 2. OBJETIVOS
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("2. Objetivos", level=1)
    doc.add_heading("2.1 Objetivo General", level=2)
    doc.add_paragraph("Aplicar técnicas de informática forense para analizar y documentar la actividad de un usuario sospechoso a través de los registros de eventos de Windows, identificando posibles huellas de actividad maliciosa.")

    doc.add_heading("2.2 Objetivos Específicos", level=2)
    objetivos = [
        "Configurar un escenario forense utilizando el Visor de Eventos (Event Viewer) de Windows.",
        "Identificar y clasificar eventos clave relacionados con actividades sospechosas en los registros de Seguridad y Sistema.",
        "Documentar detalladamente cada evento encontrado, incluyendo su ID, fecha y hora, usuario involucrado y posible interpretación forense.",
        "Analizar la correlación entre los eventos identificados para determinar patrones de comportamiento potencialmente malicioso.",
        "Presentar conclusiones fundamentadas sobre los hallazgos de la investigación.",
    ]
    for obj in objetivos:
        doc.add_paragraph(obj, style='List Bullet')

    # ═══════════════════════════════════════════════════════════════════════
    # 3. MARCO TEÓRICO
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("3. Marco Teórico", level=1)
    doc.add_heading("3.1 El Principio de Intercambio de Locard en el Entorno Digital", level=2)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("El principio de Locard, formulado por el criminólogo francés Edmond Locard, establece que cuando dos objetos entran en contacto, siempre hay una transferencia de material entre ellos. Aplicado al ámbito digital, este principio se manifiesta en cuatro dimensiones fundamentales:")

    dimensiones = [
        ("Intercambio de información:", "Cada interacción con un sistema digital genera datos. Los elementos clave incluyen archivos, logs de actividad y metadatos."),
        ("Transferencia de evidencia:", "Cada acción en el entorno digital deja huellas que pueden ser rastreadas, incluso si se intenta eliminarlas."),
        ("Persistencia de la evidencia:", "La evidencia digital puede mantenerse accesible y recuperable, incluso después de intentos de eliminación."),
        ("Análisis del intercambio:", "Los investigadores forenses analizan meticulosamente todos los rastros de intercambio digital para identificar acciones, responsables y cronología de eventos."),
    ]
    for i, (titulo, desc) in enumerate(dimensiones, 1):
        p = doc.add_paragraph()
        add_formatted_run(p, f"{i}. {titulo} ", bold=True, size=11)
        p.add_run(desc)

    doc.add_heading("3.2 El Visor de Eventos de Windows como Herramienta Forense", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("El Visor de Eventos (Event Viewer) es una herramienta integrada en Windows que registra eventos significativos del sistema operativo. Desde la perspectiva forense, los registros más relevantes son:")

    doc.add_paragraph("Registro de Seguridad: Contiene eventos de auditoría como inicios de sesión, cambios de permisos, acceso a recursos y modificaciones de cuentas de usuario.", style='List Bullet')
    doc.add_paragraph("Registro de Sistema: Contiene eventos generados por componentes del sistema operativo, incluyendo instalación de servicios, errores de drivers y cambios en la configuración del sistema.", style='List Bullet')

    doc.add_heading("3.3 Eventos Críticos para la Investigación Forense", level=2)

    event_headers = ["Event ID", "Registro", "Descripción", "Relevancia Forense"]
    event_rows = [
        ["4624", "Seguridad", "Inicio de sesión exitoso", "Identificar accesos en horarios inusuales"],
        ["4625", "Seguridad", "Inicio de sesión fallido", "Detectar intentos de acceso no autorizado"],
        ["4611", "Seguridad", "Proceso de inicio de sesión de confianza registrado", "Detectar elevación de privilegios vía UAC"],
        ["4672", "Seguridad", "Privilegios especiales asignados", "Detectar escalación de privilegios"],
        ["4697", "Seguridad", "Se instaló un servicio en el sistema", "Detectar instalación de servicios sospechosos"],
        ["4720", "Seguridad", "Cuenta de usuario creada", "Detectar creación de backdoor"],
        ["4726", "Seguridad", "Cuenta de usuario eliminada", "Detectar eliminación de evidencia"],
        ["4728", "Seguridad", "Miembro agregado a grupo global", "Detectar modificación de membresías"],
        ["4729", "Seguridad", "Miembro removido de grupo global", "Detectar eliminación de membresías"],
        ["7045", "Sistema", "Servicio instalado", "Detectar instalación de malware como servicio"],
        ["1102", "Seguridad", "Registro de auditoría borrado", "Detectar intento de eliminación de evidencia"],
    ]
    add_table(doc, event_headers, event_rows)

    # ═══════════════════════════════════════════════════════════════════════
    # 4. METODOLOGÍA
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("4. Metodología", level=1)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("Para la realización de este análisis forense se siguió la metodología forense digital compuesta por las siguientes fases (De Haro Olmo, 2020):")

    fases = [
        ("Identificación:", "Se identificaron las fuentes de evidencia digital disponibles en el sistema operativo Windows, específicamente los registros de eventos accesibles mediante el Visor de Eventos."),
        ("Preservación:", "Se documentaron los eventos mediante capturas de pantalla que garantizan la integridad visual de la evidencia en el momento de su recolección."),
        ("Análisis:", "Se examinaron detalladamente los eventos registrados, correlacionando IDs de eventos, marcas de tiempo, cuentas de usuario involucradas y descripciones para identificar patrones de actividad sospechosa."),
        ("Presentación:", "Los hallazgos se organizaron de forma estructurada en el presente documento, con evidencia visual y análisis interpretativo de cada evento."),
    ]
    for i, (titulo, desc) in enumerate(fases, 1):
        p = doc.add_paragraph()
        add_formatted_run(p, f"{i}. {titulo} ", bold=True, size=11)
        p.add_run(desc)

    doc.add_heading("4.1 Herramientas Utilizadas", level=2)
    herramientas = [
        ("Sistema operativo:", "Windows 11"),
        ("Equipo analizado:", "MELQUI"),
        ("Herramienta de análisis:", "Visor de Eventos (Event Viewer) — eventvwr.msc"),
        ("Documentación:", "Capturas de pantalla tomadas con la herramienta de recorte de Windows (Win + Shift + S)"),
        ("Configuración adicional:", "Se habilitó la auditoría de errores de inicio de sesión mediante auditpol"),
    ]
    add_key_value_table(doc, herramientas)

    doc.add_heading("4.2 Escenario de Investigación", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("Se configuró un escenario forense simulado en el cual se generaron deliberadamente eventos que replican el comportamiento típico de un atacante o usuario malintencionado:")

    escenarios = [
        "Análisis de inicios de sesión en horarios inusuales (identificación de accesos nocturnos sospechosos)",
        "Inicio de sesión con privilegios elevados de administrador",
        "Creación y eliminación de una cuenta de usuario temporal (simulación de backdoor)",
        "Instalación y eliminación de un servicio del sistema (simulación de persistencia de malware)",
        "Borrado de los registros de auditoría (simulación de técnica anti-forense)",
    ]
    for e in escenarios:
        doc.add_paragraph(e, style='List Bullet')

    # ═══════════════════════════════════════════════════════════════════════
    # 5. CONFIGURACIÓN DEL ESCENARIO FORENSE
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("5. Configuración del Escenario Forense", level=1)

    # 5.1
    doc.add_heading("5.1 Acceso al Visor de Eventos", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run('Se accedió al Visor de Eventos mediante el comando eventvwr.msc desde la ventana Ejecutar (Win + R). En la captura se puede observar simultáneamente la ventana "Ejecutar" con el comando ingresado y el Visor de Eventos ya abierto en segundo plano.')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Captura 1: Ventana Ejecutar con eventvwr.msc y Visor de Eventos abierto", bold=True, size=9, color=RGBColor(31, 78, 121))
    add_image_safe(doc, IMG_DIR / "01_abrir_event_viewer.png")

    # 5.2
    doc.add_heading("5.2 Vista general del Visor de Eventos — Árbol de navegación", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("Al expandir el árbol de navegación del panel izquierdo, se observó la estructura completa del Visor de Eventos del equipo local, incluyendo Vistas personalizadas, Registros de Windows (Aplicación, Seguridad, Instalación, Sistema, Eventos reenviados), Registros de aplicaciones y servicios, y Suscripciones.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Captura 2: Árbol de navegación completo del Visor de Eventos", bold=True, size=9, color=RGBColor(31, 78, 121))
    add_image_safe(doc, IMG_DIR / "02_visor_eventos_principal.png")

    # 5.3
    doc.add_heading("5.3 Registro de Seguridad", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("Se navegó a Registros de Windows > Seguridad para revisar los eventos de auditoría del sistema. El registro contiene un total de 21.312 eventos al momento del análisis. Los eventos predominantes corresponden al Event ID 4662 (Se realizó una operación en un objeto), categoría \"Other Object Access Events\", con nivel de auditoría correcta.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Captura 3: Registro de Seguridad mostrando 21.312 eventos, predominando Event ID 4662", bold=True, size=9, color=RGBColor(31, 78, 121))
    add_image_safe(doc, IMG_DIR / "03_registro_seguridad.png")

    # Tabla de sesiones
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("Al filtrar los eventos de inicio y cierre de sesión, se identificaron múltiples eventos 4624 (Logon) y 4634 (Logoff) correspondientes a la actividad del usuario en los días 13, 14 y 15 de febrero de 2026:")

    logon_headers = ["Fecha y hora", "Event ID", "Tipo", "Observación"]
    logon_rows = [
        ["14/2/2026 21:30:14", "4624", "Logon", "⚠️ Horario nocturno inusual"],
        ["14/2/2026 21:35:47", "4624", "Logon", "⚠️ Horario nocturno inusual"],
        ["14/2/2026 21:20:54", "4624", "Logon", "⚠️ Horario nocturno inusual"],
        ["14/2/2026 23:19:25", "4634", "Logoff", "⚠️ Cierre de sesión en madrugada"],
        ["14/2/2026 18:00:23", "4624", "Logon", "Horario fuera de jornada laboral"],
        ["14/2/2026 18:29:12", "4624", "Logon", "Horario fuera de jornada laboral"],
        ["14/2/2026 18:16:38", "4624", "Logon", "Horario fuera de jornada laboral"],
        ["14/2/2026 12:26:41", "4624", "Logon", "Horario laboral normal"],
        ["14/2/2026 12:46:49", "4624", "Logon", "Horario laboral normal"],
        ["15/2/2026 12:51:01", "4624", "Logon", "Horario laboral normal"],
        ["15/2/2026 12:46:37", "4624", "Logon", "Horario laboral normal"],
        ["15/2/2026 14:36:01", "4624", "Logon", "Horario laboral normal"],
        ["15/2/2026 12:42:19", "4624", "Logon", "Horario laboral normal"],
        ["14/2/2026 15:53:02", "4624", "Logon", "Horario laboral normal"],
        ["13/2/2026 17:03:58", "4624", "Logon", "Horario laboral normal"],
        ["13/2/2026 16:44:20", "4624", "Logon", "Horario laboral normal"],
        ["13/2/2026 21:59:13", "4634", "Logoff", "⚠️ Cierre de sesión nocturno"],
    ]
    add_table(doc, logon_headers, logon_rows)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("Del análisis de los registros de inicio y cierre de sesión se identificó un patrón relevante: existen múltiples accesos en horarios nocturnos (después de las 21:00 horas), lo cual puede ser indicador de actividad sospechosa.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Captura 5: Registro de Seguridad filtrado mostrando eventos 4624 (Logon) y 4634 (Logoff)", bold=True, size=9, color=RGBColor(31, 78, 121))
    add_image_safe(doc, IMG_DIR / "05_eventos_logon.png")

    # 5.4
    doc.add_heading("5.4 Registro de Sistema", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("Se navegó a Registros de Windows > Sistema para revisar los eventos del sistema operativo. El registro contiene un total de 28.280 eventos. Se observaron eventos predominantes de tipo BTHUSB (Event ID 12, Advertencia) y DistributedCOM (Event ID 10016, Advertencia).")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Captura 4: Registro de Sistema mostrando 28.280 eventos", bold=True, size=9, color=RGBColor(31, 78, 121))
    add_image_safe(doc, IMG_DIR / "04_registro_sistema.png")

    # ═══════════════════════════════════════════════════════════════════════
    # 6. RESUMEN DE HALLAZGOS
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("6. Resumen de Hallazgos", level=1)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("Durante el análisis forense de los registros de eventos del equipo MELQUI (21.365 eventos en Seguridad, 28.280 eventos en Sistema) se identificaron los siguientes eventos sospechosos:")

    resumen_headers = ["#", "Event ID", "Registro", "Fecha y hora", "Categoría", "Interpretación"]
    resumen_rows = [
        ["1", "4624", "Seguridad", "14/2/2026 21:20:54", "Logon — Nocturno", "Primer acceso nocturno fuera de jornada"],
        ["2", "4624", "Seguridad", "14/2/2026 21:30:14", "Logon — Nocturno", "Segundo acceso nocturno en 10 min"],
        ["3", "4624", "Seguridad", "14/2/2026 21:35:47", "Logon — Nocturno", "Tercer acceso nocturno"],
        ["4", "4634", "Seguridad", "14/2/2026 23:19:25", "Logoff — Madrugada", "Cierre de sesión en madrugada"],
        ["5", "4611", "Seguridad", "15/2/2026 17:02:38", "Security System Ext.", "Registro proceso confianza (UAC)"],
        ["6", "4720", "Seguridad", "15/2/2026 20:57:16", "User Account Mgmt.", "Creación cuenta UsuarioPrueba"],
        ["7", "4728", "Seguridad", "15/2/2026 20:57:16", "Security Group Mgmt.", "UsuarioPrueba → grupo Ninguno"],
        ["8", "4726", "Seguridad", "15/2/2026 20:58:10", "User Account Mgmt.", "Eliminación cuenta UsuarioPrueba"],
        ["9", "4729", "Seguridad", "15/2/2026 20:58:10", "Security Group Mgmt.", "UsuarioPrueba removido de grupo"],
        ["10", "4697", "Seguridad", "15/2/2026 21:13:32", "Security System Ext.", "Instalación ServicioPrueba (cmd.exe)"],
        ["11", "1102", "Seguridad", "15/2/2026 21:23:32", "Log clear", "Borrado intencional de logs"],
    ]
    add_table(doc, resumen_headers, resumen_rows)

    # ═══════════════════════════════════════════════════════════════════════
    # 7. DESCRIPCIÓN DETALLADA DE EVENTOS SOSPECHOSOS
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("7. Descripción Detallada de Eventos Sospechosos", level=1)

    # ─── 7.1 Event 4624 ──────────────────────────────────────────────────
    doc.add_heading("7.1 Evento: Inicios de Sesión en Horarios Inusuales (Event ID 4624)", level=2)

    p = doc.add_paragraph()
    add_formatted_run(p, "Tipo de evento: ", bold=True)
    p.add_run("Auditoría correcta — Inicio de sesión (Logon)")
    p = doc.add_paragraph()
    add_formatted_run(p, "Registro: ", bold=True)
    p.add_run("Seguridad")
    p = doc.add_paragraph()
    add_formatted_run(p, "Nivel de sospecha: ", bold=True)
    add_formatted_run(p, "Alto", bold=True, color=RGBColor(200, 0, 0))

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("Durante la revisión de los registros de seguridad del equipo MELQUI, se identificó un patrón de múltiples inicios de sesión exitosos en horarios nocturnos que no corresponden a una jornada laboral convencional (8:00 – 18:00 horas). Se detectaron tres sesiones en un intervalo de 15 minutos durante la noche del 14 de febrero de 2026, seguidas de un cierre de sesión casi dos horas después.")

    add_key_value_table(doc, [
        ("ID del evento", "4624"),
        ("Acceso sospechoso #1", "14/2/2026 21:20:54"),
        ("Acceso sospechoso #2", "14/2/2026 21:30:14"),
        ("Acceso sospechoso #3", "14/2/2026 21:35:47"),
        ("Cierre de sesión", "14/2/2026 23:19:25 (Event ID 4634)"),
        ("Duración estimada", "Aproximadamente 2 horas"),
        ("Equipo", "MELQUI"),
    ])

    doc.add_paragraph("")
    p = doc.add_paragraph()
    add_formatted_run(p, "Posible interpretación forense:", bold=True)

    interp_4624 = [
        "Se registraron 3 inicios de sesión entre las 21:20 y las 21:35, un intervalo de solo 15 minutos.",
        "La sesión se mantuvo activa hasta las 23:19, lo cual indica operación del equipo durante casi 2 horas en horario nocturno.",
        "En una investigación real, se correlacionaría esta actividad con otros indicadores: programas ejecutados, archivos accedidos, conexiones de red.",
    ]
    for item in interp_4624:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph("")
    p = doc.add_paragraph()
    add_formatted_run(p, "Comparación de patrones de acceso:", bold=True)

    patrones_headers = ["Período", "Cantidad de accesos (4624)", "Evaluación"]
    patrones_rows = [
        ["13/2 - Jornada diurna", "2 eventos", "Normal"],
        ["14/2 - Jornada diurna", "8 eventos", "Normal / Alto"],
        ["14/2 - Horario nocturno", "6 eventos", "⚠️ Sospechoso"],
        ["15/2 - Jornada diurna", "5 eventos", "Normal"],
    ]
    add_table(doc, patrones_headers, patrones_rows)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Captura 6: Eventos 4624 con horarios sospechosos", bold=True, size=9, color=RGBColor(31, 78, 121))
    add_image_safe(doc, IMG_DIR / "06_eventos_logon_sospechosos.png")

    # ─── 7.2 Event 4611/4672 ─────────────────────────────────────────────
    doc.add_heading("7.2 Evento: Escalación de Privilegios (Event ID 4611 / 4672)", level=2)

    p = doc.add_paragraph()
    add_formatted_run(p, "Tipo de evento: ", bold=True)
    p.add_run("Auditoría correcta — Security System Extension")
    p = doc.add_paragraph()
    add_formatted_run(p, "Registro: ", bold=True)
    p.add_run("Seguridad")
    p = doc.add_paragraph()
    add_formatted_run(p, "Nivel de sospecha: ", bold=True)
    add_formatted_run(p, "Alto", bold=True, color=RGBColor(200, 0, 0))

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("Al abrir una sesión de PowerShell como Administrador, el sistema operativo invocó el proceso ConsentUI (la ventana de Control de Cuentas de Usuario — UAC) para autorizar la elevación de privilegios. Windows registró este evento como Event ID 4611 (\"Se registró un proceso de inicio de sesión de confianza\"). Adicionalmente, se observó el Event ID 4672 (Special Logon), que confirma la asignación efectiva de los privilegios especiales de administrador.")

    add_key_value_table(doc, [
        ("ID del evento principal", "4611"),
        ("Fecha y hora", "15/2/2026 17:02:38"),
        ("Categoría de tarea", "Security System Extension"),
        ("Nombre de cuenta", "MELQUI$"),
        ("Nombre de proceso de inicio de sesión", "ConsentUI"),
        ("Equipo", "MELQUI"),
        ("Evento complementario", "4672 (Special Logon) — 15/2/2026 20:54:59"),
    ])

    doc.add_paragraph("")
    p = doc.add_paragraph()
    add_formatted_run(p, "Posible interpretación forense:", bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("La elevación de privilegios mediante UAC (ConsentUI) es un paso clave en la cadena de ataque. Un atacante con privilegios de administrador puede instalar software malicioso, crear cuentas de usuario para acceso futuro, modificar configuraciones de seguridad y borrar registros de auditoría.")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Captura 7: Event ID 4611 — ConsentUI al abrir PowerShell como Administrador", bold=True, size=9, color=RGBColor(31, 78, 121))
    add_image_safe(doc, IMG_DIR / "07_evento_4611_detalle.png")

    # ─── 7.3 Event 4720 + 4728 ──────────────────────────────────────────
    doc.add_heading("7.3 Evento: Creación de Cuenta de Usuario (Event ID 4720)", level=2)

    p = doc.add_paragraph()
    add_formatted_run(p, "Tipo de evento: ", bold=True)
    p.add_run("Auditoría correcta — User Account Management")
    p = doc.add_paragraph()
    add_formatted_run(p, "Registro: ", bold=True)
    p.add_run("Seguridad")
    p = doc.add_paragraph()
    add_formatted_run(p, "Nivel de sospecha: ", bold=True)
    add_formatted_run(p, "Alto", bold=True, color=RGBColor(200, 0, 0))

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("Se detectó la creación de una nueva cuenta de usuario en el sistema. La cuenta fue creada mediante el comando ")
    add_formatted_run(p, "net user UsuarioPrueba Password123! /add", italic=True)
    p.add_run(" ejecutado desde PowerShell con privilegios de administrador. Adicionalmente, Windows generó automáticamente el Event ID 4728 (\"Se agregó un miembro a un grupo global con seguridad habilitada\"), registrando que UsuarioPrueba fue añadido al grupo \"Ninguno\" del equipo MELQUI.")

    add_key_value_table(doc, [
        ("ID del evento", "4720"),
        ("Fecha y hora", "15/2/2026 20:57:16"),
        ("Categoría de tarea", "User Account Management"),
        ("Cuenta creada", "UsuarioPrueba"),
        ("Dominio de cuenta creada", "MELQUI"),
        ("Creada por (Sujeto)", "melqui.romero (SENA\\melqui.romero)"),
        ("Id. de inicio de sesión", "0xF2F83D8"),
        ("Equipo", "MELQUI"),
    ])

    doc.add_paragraph("")
    p = doc.add_paragraph()
    add_formatted_run(p, "Evento asociado 4728:", bold=True)

    add_key_value_table(doc, [
        ("ID del evento", "4728"),
        ("Fecha y hora", "15/2/2026 20:57:16"),
        ("Miembro agregado", "MELQUI\\UsuarioPrueba"),
        ("Grupo destino", "Ninguno (MELQUI\\Ninguno)"),
        ("Acción realizada por", "melqui.romero (SENA)"),
    ])

    doc.add_paragraph("")
    p = doc.add_paragraph()
    add_formatted_run(p, "Posible interpretación forense:", bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("La creación de una cuenta de usuario sin justificación legítima es una técnica de persistencia empleada por atacantes. Permite mantener acceso al sistema, operar bajo una cuenta diferente y escalar privilegios. El evento complementario 4728 genera evidencia forense adicional que el atacante podría no anticipar.")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Captura 8: Terminal con comando net user y popup del Event ID 4720", bold=True, size=9, color=RGBColor(31, 78, 121))
    add_image_safe(doc, IMG_DIR / "08_crear_usuario_cmd.png")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Captura 10: Detalles del Event ID 4728 — Miembro agregado a grupo", bold=True, size=9, color=RGBColor(31, 78, 121))
    add_image_safe(doc, IMG_DIR / "10_evento_4728_detalle.png")

    # ─── 7.4 Event 4726 + 4729 ──────────────────────────────────────────
    doc.add_heading("7.4 Evento: Eliminación de Cuenta de Usuario (Event ID 4726)", level=2)

    p = doc.add_paragraph()
    add_formatted_run(p, "Tipo de evento: ", bold=True)
    p.add_run("Auditoría correcta — User Account Management")
    p = doc.add_paragraph()
    add_formatted_run(p, "Registro: ", bold=True)
    p.add_run("Seguridad")
    p = doc.add_paragraph()
    add_formatted_run(p, "Nivel de sospecha: ", bold=True)
    add_formatted_run(p, "Alto", bold=True, color=RGBColor(200, 0, 0))

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("Se detectó la eliminación de la cuenta de usuario creada previamente mediante ")
    add_formatted_run(p, "net user UsuarioPrueba /delete", italic=True)
    p.add_run(", apenas 54 segundos después de la creación (20:57:16 → 20:58:10). De forma análoga, la eliminación generó automáticamente el Event ID 4729 (remoción del grupo \"Ninguno\").")

    add_key_value_table(doc, [
        ("ID del evento", "4726"),
        ("Fecha y hora", "15/2/2026 20:58:10"),
        ("Categoría de tarea", "User Account Management"),
        ("Cuenta eliminada", "UsuarioPrueba"),
        ("Eliminada por (Sujeto)", "melqui.romero (SENA\\melqui.romero)"),
        ("Equipo", "MELQUI"),
        ("Tiempo entre creación y eliminación", "54 segundos"),
    ])

    doc.add_paragraph("")
    p = doc.add_paragraph()
    add_formatted_run(p, "Evento asociado 4729:", bold=True)

    add_key_value_table(doc, [
        ("ID del evento", "4729"),
        ("Fecha y hora", "15/2/2026 20:58:10"),
        ("Miembro removido", "MELQUI\\UsuarioPrueba"),
        ("Grupo", "Ninguno (MELQUI\\Ninguno)"),
        ("Acción realizada por", "melqui.romero (SENA)"),
    ])

    doc.add_paragraph("")
    p = doc.add_paragraph()
    add_formatted_run(p, "Posible interpretación forense:", bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("La eliminación de una cuenta recién creada es un indicador clásico de actividad anti-forense. Sin embargo, como demuestra el principio de Locard, esta acción deja múltiples rastros: el evento 4726, el evento 4729 y la correlación temporal con el evento 4720.")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Captura 9: Terminal mostrando creación y eliminación de usuario con lista de eventos", bold=True, size=9, color=RGBColor(31, 78, 121))
    add_image_safe(doc, IMG_DIR / "09_eliminar_usuario_cmd.png")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Captura 11: Detalles del Event ID 4729 — Miembro removido de grupo", bold=True, size=9, color=RGBColor(31, 78, 121))
    add_image_safe(doc, IMG_DIR / "11_evento_4729_detalle.png")

    # ─── 7.5 Event 4697 ─────────────────────────────────────────────────
    doc.add_heading("7.5 Evento: Instalación de Servicio del Sistema (Event ID 4697)", level=2)

    p = doc.add_paragraph()
    add_formatted_run(p, "Tipo de evento: ", bold=True)
    p.add_run("Auditoría correcta — Security System Extension")
    p = doc.add_paragraph()
    add_formatted_run(p, "Registro: ", bold=True)
    p.add_run("Seguridad")
    p = doc.add_paragraph()
    add_formatted_run(p, "Nivel de sospecha: ", bold=True)
    add_formatted_run(p, "Alto", bold=True, color=RGBColor(200, 0, 0))

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("Se detectó la instalación de un nuevo servicio en el sistema mediante ")
    add_formatted_run(p, 'sc.exe create ServicioPrueba binPath= "C:\\Windows\\System32\\cmd.exe" start= demand', italic=True)
    p.add_run(". Windows registró esta instalación como Event ID 4697 en el registro de Seguridad.")

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    add_formatted_run(p, "Nota técnica: ", bold=True, italic=True)
    p.add_run("La instalación de servicios puede registrarse como Event ID 7045 en Sistema o como Event ID 4697 en Seguridad, dependiendo de la configuración de auditoría. En este equipo, el registro se capturó en Seguridad como 4697.")

    add_key_value_table(doc, [
        ("ID del evento", "4697"),
        ("Fecha y hora", "15/2/2026 21:13:32"),
        ("Categoría de tarea", "Security System Extension"),
        ("Nombre del servicio", "ServicioPrueba"),
        ("Ruta del archivo del servicio", "C:\\Windows\\System32\\cmd.exe"),
        ("Tipo de servicio", "0x10 (propio)"),
        ("Tipo de inicio del servicio", "3 (demand / bajo demanda)"),
        ("Cuenta de servicio", "LocalSystem"),
        ("Instalado por (Sujeto)", "melqui.romero (SENA\\melqui.romero)"),
        ("Id. de inicio de sesión", "0xF2F83D8"),
        ("Equipo", "MELQUI"),
    ])

    doc.add_paragraph("")
    p = doc.add_paragraph()
    add_formatted_run(p, "Posible interpretación forense:", bold=True)

    interp_4697 = [
        "cmd.exe es el intérprete de comandos de Windows, no un servicio legítimo",
        "Un atacante podría usar este mecanismo para ejecutar comandos arbitrarios con privilegios de SYSTEM",
        "La cuenta de servicio es LocalSystem, el nivel más alto de privilegios en Windows",
        'El tipo de inicio "demand" permite al atacante activar su payload en el momento deseado',
    ]
    for item in interp_4697:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Captura 12: Terminal mostrando creación y eliminación de servicio", bold=True, size=9, color=RGBColor(31, 78, 121))
    add_image_safe(doc, IMG_DIR / "12_crear_servicio_cmd.png")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Captura 13: Detalles del Event ID 4697 — Servicio instalado", bold=True, size=9, color=RGBColor(31, 78, 121))
    add_image_safe(doc, IMG_DIR / "13_evento_4697_detalle.png")

    # ─── 7.6 Event 1102 ─────────────────────────────────────────────────
    doc.add_heading("7.6 Evento: Borrado del Registro de Auditoría (Event ID 1102)", level=2)

    p = doc.add_paragraph()
    add_formatted_run(p, "Tipo de evento: ", bold=True)
    p.add_run("Auditoría correcta — Log clear")
    p = doc.add_paragraph()
    add_formatted_run(p, "Registro: ", bold=True)
    p.add_run("Seguridad")
    p = doc.add_paragraph()
    add_formatted_run(p, "Nivel de sospecha: ", bold=True)
    add_formatted_run(p, "CRÍTICO", bold=True, color=RGBColor(200, 0, 0))

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("Se detectó que los registros de auditoría de seguridad fueron borrados intencionalmente mediante el comando ")
    add_formatted_run(p, "wevtutil cl Security", italic=True)
    p.add_run(". Tras la ejecución, el registro de Seguridad quedó con únicamente 1 evento: el propio Event ID 1102 que documenta el borrado. Esto confirma la destrucción exitosa de todos los registros anteriores.")

    add_key_value_table(doc, [
        ("ID del evento", "1102"),
        ("Fecha y hora", "15/2/2026 21:23:32"),
        ("Categoría de tarea", "Log clear"),
        ("Origen", "Eventlog"),
        ("Usuario que borró los logs", "melqui.romero (SENA\\melqui.romero)"),
        ("Id. de inicio de sesión", "0xF2F83D8"),
        ("Equipo", "MELQUI"),
        ("Eventos restantes después del borrado", "1 (solo el propio 1102)"),
    ])

    doc.add_paragraph("")
    p = doc.add_paragraph()
    add_formatted_run(p, "Posible interpretación forense:", bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run('El borrado de registros de auditoría es una de las acciones más sospechosas en un análisis forense. Paradójicamente, el propio sistema genera un evento (1102) que registra esta acción — un ejemplo perfecto del principio de Locard: "incluso el intento de eliminar los rastros deja un rastro". El Id. de inicio de sesión (0xF2F83D8) coincide con el de los eventos 4720, 4726 y 4697, confirmando que todas estas acciones fueron realizadas dentro de la misma sesión de usuario.')

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_run(p, "Captura 14: Event ID 1102 — Registro de auditoría borrado", bold=True, size=9, color=RGBColor(31, 78, 121))
    add_image_safe(doc, IMG_DIR / "14_evento_1102_detalle.png")

    # ═══════════════════════════════════════════════════════════════════════
    # 8. LÍNEA DE TIEMPO DEL INCIDENTE
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("8. Línea de Tiempo del Incidente", level=1)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.add_run("A continuación, se presenta la reconstrucción cronológica del incidente basada en los eventos analizados:")

    timeline_headers = ["#", "Fecha y hora", "Evento", "Interpretación"]
    timeline_rows = [
        ["1", "14/2/2026 21:20", "Inicio de sesión nocturno (4624)", "Acceso sospechoso fuera de jornada"],
        ["2", "14/2/2026 21:30", "Segundo inicio nocturno (4624)", "Patrón de accesos repetidos"],
        ["3", "14/2/2026 21:35", "Tercer inicio nocturno (4624)", "Actividad sostenida fuera de horario"],
        ["4", "15/2/2026 17:02", "Elevación de privilegios UAC (4611)", "Escalación mediante ConsentUI"],
        ["5", "15/2/2026 20:57", "Cuenta creada (4720) + Grupo (4728)", "Creación de backdoor"],
        ["6", "15/2/2026 20:58", "Cuenta eliminada (4726) + Grupo (4729)", "Destrucción evidencia (54 seg.)"],
        ["7", "15/2/2026 21:13", "Servicio instalado (4697)", "Servicio sospechoso → cmd.exe"],
        ["8", "15/2/2026 21:23", "Logs borrados (1102)", "Técnica anti-forense"],
        ["9", "14/2/2026 23:19", "Cierre de sesión (4634)", "Fin de actividad sospechosa"],
    ]
    add_table(doc, timeline_headers, timeline_rows, header_color="8B0000")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    add_formatted_run(p, "Esta secuencia de eventos sigue un patrón de ataque conocido como kill chain:", bold=True)

    kill_chain = [
        "Acceso inicial → Inicios de sesión en horarios nocturnos inusuales",
        "Escalación de privilegios → Obtención de privilegios de administrador",
        "Persistencia → Creación de cuentas backdoor e instalación de servicios",
        "Limpieza de huellas → Eliminación de cuentas temporales y borrado de logs",
    ]
    for i, item in enumerate(kill_chain, 1):
        p = doc.add_paragraph()
        add_formatted_run(p, f"{i}. ", bold=True)
        p.add_run(item)

    # ═══════════════════════════════════════════════════════════════════════
    # 9. CONCLUSIONES
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("9. Conclusiones", level=1)

    conclusiones = [
        ("Aplicación del principio de Locard:", "El análisis realizado demuestra que el principio de intercambio de Locard se aplica plenamente en el entorno digital. Cada acción ejecutada en el sistema —desde un inicio de sesión hasta el borrado de registros— dejó un rastro identificable en los logs del sistema operativo. El equipo MELQUI registró más de 21.000 eventos de seguridad y 28.000 eventos de sistema. Incluso las técnicas anti-forense como el borrado de logs generan su propio evento (ID 1102), confirmando que \"todo contacto deja un rastro\"."),
        ("Valor forense del Visor de Eventos:", "El Event Viewer de Windows demostró ser una herramienta fundamental para la investigación forense de primer nivel. Los registros de Seguridad y Sistema proporcionan información detallada y cronológica de las acciones realizadas en el equipo."),
        ("Detección de patrones de acceso sospechoso:", "Se identificaron múltiples inicios de sesión (Event ID 4624) en horarios nocturnos inusuales, particularmente el 14/2/2026 entre las 21:20 y las 21:35, con un cierre de sesión a las 23:19. Este patrón constituye un indicador de compromiso (IoC)."),
        ("Identificación de técnicas de ataque:", "Las pruebas de simulación forense evidenciaron que un atacante puede elevar privilegios mediante UAC (4611), crear cuentas de backdoor (4720/4728), instalar servicios maliciosos (4697) y borrar registros de auditoría (1102). Este patrón corresponde al modelo de la kill chain (De Haro Olmo, 2020). Además, cada acción genera múltiples eventos complementarios (e.g., 4720+4728, 4726+4729), multiplicando las fuentes de evidencia."),
        ("Importancia de la cadena de custodia:", "La documentación rigurosa de cada evento con capturas de pantalla, timestamps, IDs de eventos y descripción técnica es esencial para mantener la validez probatoria de la evidencia digital (Luna Felipez, s.f.)."),
        ("Necesidad de monitoreo continuo:", "Los hallazgos evidencian la importancia de implementar sistemas de monitoreo y alertas que detecten en tiempo real patrones sospechosos como accesos en horarios inusuales, escalación de privilegios, creación de cuentas no autorizadas o borrado de logs."),
    ]
    for i, (titulo, desc) in enumerate(conclusiones, 1):
        p = doc.add_paragraph()
        add_formatted_run(p, f"{i}. {titulo} ", bold=True)
        p.add_run(desc)

    # ═══════════════════════════════════════════════════════════════════════
    # 10. BIBLIOGRAFÍA
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("10. Bibliografía", level=1)

    refs = [
        "De Haro Olmo, F. J. (2020). Crimen, cibercrimen y análisis forense informático. I.E.S. Celia Viñas. Recuperado de https://iescelia.org/ciberseguridad/ceceti-afi-00",
        "Luna Felipez, J. P. (s.f.). Perito Informático Judicial Forense. Universidad Nacional \"Siglo XX\", Llallagua, Bolivia.",
        "Areito, G. (2008). Seguridad de la Información: Redes, Informática y Sistemas de Información.",
        "Microsoft. (2024). Windows Security auditing. Microsoft Learn. Recuperado de https://learn.microsoft.com/en-us/windows/security/threat-protection/auditing/",
        "NIST. (2012). Guide to Computer Security Log Management (SP 800-92). National Institute of Standards and Technology.",
    ]
    for ref in refs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(ref)
        p.paragraph_format.space_after = Pt(4)

    # ─── Guardar ────────────────────────────────────────────────────────
    doc.save(str(OUTPUT_FILE))
    print(f"\n✅ Documento generado exitosamente:")
    print(f"   {OUTPUT_FILE}")
    print(f"   Tamaño: {OUTPUT_FILE.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build_document()
