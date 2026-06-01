"""
Script para generar los 3 archivos .docx de la Actividad Evaluativa Eje 4.
Cada archivo representa una perspectiva diferente sobre el caso Diego Gómez.

Uso: python generate_docx.py
Requisito: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ENTREGA_DIR = os.path.join(BASE_DIR, "entrega")
os.makedirs(ENTREGA_DIR, exist_ok=True)


def set_normal_style(doc):
    """Configura el estilo normal del documento."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15


def add_title_page(doc, estudiante, perspectiva):
    """Agrega página de título al documento."""
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Actividad Evaluativa – Eje 4')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Análisis del Caso: Estudiante podría ir a cárcel\npor divulgar tesis en una plataforma')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(51, 51, 51)

    doc.add_paragraph()

    materia = doc.add_paragraph()
    materia.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = materia.add_run('Aspectos Legales de la Seguridad Informática')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'Estudiante: {estudiante}')
    run.font.size = Pt(12)

    enfoque = doc.add_paragraph()
    enfoque.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = enfoque.add_run(f'Enfoque: {perspectiva}')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    uni = doc.add_paragraph()
    uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = uni.add_run('Universidad Areandina')
    run.font.size = Pt(12)

    doc.add_paragraph()

    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fecha.add_run('2026')
    run.font.size = Pt(12)

    doc.add_page_break()


def add_heading_styled(doc, text, level=1):
    """Agrega un heading con estilo personalizado."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return heading


def add_bold_paragraph(doc, bold_text, normal_text=""):
    """Agrega un párrafo con texto en negrita seguido de texto normal."""
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p


def add_bullet(doc, text, bold_prefix=""):
    """Agrega un ítem de lista con viñeta."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_table(doc, headers, rows):
    """Agrega una tabla con estilo."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # Espacio después de la tabla
    return table


# ==============================================================================
# ARCHIVO 1 - MELQUI (Análisis jurídico-técnico)
# ==============================================================================
def generate_melqui():
    doc = Document()
    set_normal_style(doc)
    add_title_page(doc, "Melqui Romero", "Análisis jurídico-técnico del caso")

    # --- Contexto ---
    add_heading_styled(doc, "Contexto del Caso", level=1)
    doc.add_paragraph(
        'El artículo de El Espectador relata el caso de Diego Gómez, un estudiante de biología '
        'de la Universidad del Quindío, quien en 2011 subió a la plataforma Scribd una tesis de '
        'maestría de la Universidad Nacional de Colombia que había encontrado en internet. Su '
        'intención era facilitar la consulta del documento a otros investigadores y compañeros que, '
        'al igual que él, enfrentaban dificultades para acceder a bases de datos académicas. Dos '
        'años después, en 2013, el autor de la tesis lo denunció penalmente por violación a los '
        'derechos patrimoniales de autor, enfrentándose a una posible condena de 4 a 8 años de prisión.'
    )

    # --- Pregunta 1 ---
    add_heading_styled(doc, "1. Posibles delitos en los que incurrió el ciudadano", level=1)

    add_heading_styled(doc, "a) Violación a los Derechos Patrimoniales de Autor y Derechos Conexos", level=2)
    doc.add_paragraph(
        'Este es el delito principal por el cual fue denunciado. La conducta consiste en reproducir, '
        'distribuir o comunicar públicamente una obra protegida por derechos de autor sin la autorización '
        'del titular. Al subir la tesis a Scribd, Gómez realizó una reproducción digital no autorizada '
        'y la puso a disposición del público a través de internet, lo que constituye una forma de '
        'comunicación pública y distribución de la obra.'
    )

    add_heading_styled(doc, "b) Reproducción ilícita de obra ajena", level=2)
    doc.add_paragraph(
        'Al realizar una copia digital de la tesis y subirla a una plataforma de acceso público, se '
        'configuró la reproducción no autorizada de una obra protegida, independientemente de que la '
        'intención fuera académica y sin ánimo de lucro.'
    )
    doc.add_paragraph(
        'Es importante señalar que, aunque la intención del estudiante no fue lucrativa sino académica, '
        'la legislación colombiana vigente al momento de los hechos no contemplaba excepciones claras '
        'para el uso educativo o investigativo sin autorización del autor, lo que derivó en la viabilidad '
        'de la acción penal.'
    )

    # --- Pregunta 2 ---
    add_heading_styled(doc, "2. Leyes y artículos infringidos", level=1)
    doc.add_paragraph('El estudiante habría infringido la siguiente normatividad colombiana:')

    add_table(doc,
        ["Norma", "Artículo(s)", "Descripción"],
        [
            ["Ley 23 de 1982\n(Ley de Derechos de Autor)", "Arts. 2, 3, 12, 72",
             "Establece la protección de las obras literarias, científicas y artísticas. "
             "Prohíbe la reproducción, distribución y comunicación pública sin autorización del titular."],
            ["Código Penal\n(Ley 599 de 2000)", "Art. 271",
             "Tipifica el delito de \"Violación a los derechos patrimoniales de autor y derechos conexos\". "
             "Sanciona a quien, sin autorización, reproduzca, distribuya o comunique públicamente una obra protegida."],
            ["Ley 1032 de 2006", "Modifica Art. 271",
             "Reformó el artículo 271 del Código Penal, endureciendo las penas y eliminando "
             "la necesidad de demostrar ánimo de lucro."],
            ["Decisión Andina 351/1993", "Arts. 4, 13, 52",
             "Régimen Común sobre Derecho de Autor de la Comunidad Andina, aplicable en Colombia."],
            ["Ley 44 de 1993", "Arts. 51, 52",
             "Modifica y complementa la Ley 23 de 1982 en materia de derechos morales y patrimoniales."],
        ]
    )

    # --- Pregunta 3 ---
    add_heading_styled(doc, "3. Penas que tendría que asumir de ser declarado culpable", level=1)

    add_heading_styled(doc, "Pena principal:", level=2)
    add_bullet(doc, "Prisión de cuatro (4) a ocho (8) años ", "• ")
    doc.add_paragraph('por el delito de violación a los derechos patrimoniales de autor y derechos conexos.')

    add_heading_styled(doc, "Penas accesorias:", level=2)
    add_bullet(doc, "Multa de 26.66 a 1.000 salarios mínimos legales mensuales vigentes (SMLMV).")
    add_bullet(doc, "Inhabilitación para el ejercicio de derechos y funciones públicas por un período equivalente al de la pena de prisión.")

    add_heading_styled(doc, "Consecuencias adicionales:", level=2)
    add_bullet(doc, "Antecedentes penales que afectarían su vida profesional y académica.")
    add_bullet(doc, "Reparación civil al titular de los derechos de autor por los perjuicios causados.")
    add_bullet(doc, "Decomiso de los equipos y medios utilizados para la comisión del delito.")

    doc.add_paragraph(
        'Es relevante destacar que estas penas resultan desproporcionadas frente a la conducta realizada: '
        'un estudiante que compartió un documento académico sin ánimo de lucro con fines puramente investigativos. '
        'De hecho, el caso terminó con la absolución de Diego Gómez en mayo de 2017, ratificada por el Tribunal '
        'Superior de Bogotá en diciembre del mismo año.'
    )

    # --- Pregunta 4 ---
    add_heading_styled(doc, "4. Recomendaciones como experto en seguridad", level=1)

    add_heading_styled(doc, "Recomendaciones para los ciudadanos y la comunidad académica:", level=2)
    add_bullet(doc, "Verificar los términos de uso y licencia de cualquier documento antes de compartirlo "
               "o subirlo a plataformas digitales.", "1. ")
    add_bullet(doc, "Solicitar autorización expresa al autor o titular de los derechos antes de reproducir "
               "o distribuir una obra protegida.", "2. ")
    add_bullet(doc, "Utilizar repositorios institucionales autorizados como los que ofrecen las universidades "
               "colombianas.", "3. ")
    add_bullet(doc, "Capacitarse en propiedad intelectual y derechos de autor como parte de la formación profesional.", "4. ")
    add_bullet(doc, "Hacer uso del derecho de cita (Art. 31 de la Ley 23 de 1982), que permite la utilización "
               "de fragmentos con fines de crítica, enseñanza o investigación.", "5. ")

    add_heading_styled(doc, "Recomendaciones para el Estado colombiano:", level=2)
    add_bullet(doc, "Reformar la legislación de derechos de autor para incluir excepciones claras para el uso "
               "educativo, investigativo y científico sin ánimo de lucro.", "1. ")
    add_bullet(doc, "Implementar la vía administrativa o civil como mecanismo preferente para resolver "
               "conflictos de propiedad intelectual en el ámbito académico.", "2. ")
    add_bullet(doc, "Promover políticas de acceso abierto (Open Access) a la producción científica "
               "financiada con recursos públicos.", "3. ")
    add_bullet(doc, "Crear programas de formación ciudadana sobre derechos de autor y propiedad intelectual "
               "en el ámbito digital.", "4. ")
    add_bullet(doc, "Fortalecer los repositorios digitales institucionales de las universidades.", "5. ")

    # --- Referencias ---
    doc.add_page_break()
    add_heading_styled(doc, "Referencias", level=1)
    refs = [
        "Ley 23 de 1982 – Sobre Derechos de Autor (Colombia).",
        "Ley 599 de 2000 – Código Penal Colombiano, Art. 271.",
        "Ley 1032 de 2006 – Modifica el Art. 271 del Código Penal.",
        "Decisión Andina 351 de 1993 – Régimen Común sobre Derecho de Autor.",
        "Ley 44 de 1993 – Complementa la Ley 23 de 1982.",
        "Ley 1915 de 2018 – Modernización de Derechos de Autor en Colombia.",
        "Fundación Karisma – Campaña #CompartirNoEsDelito.",
        'El Espectador (2014). "Estudiante podría ir a cárcel por divulgar tesis en una plataforma".',
    ]
    for ref in refs:
        add_bullet(doc, ref)

    path = os.path.join(ENTREGA_DIR, "Solucion_Eje4_Melqui.docx")
    doc.save(path)
    print(f"✅ Generado: {path}")


# ==============================================================================
# ARCHIVO 2 - VICTOR (Análisis crítico desde derechos digitales)
# ==============================================================================
def generate_victor():
    doc = Document()
    set_normal_style(doc)
    add_title_page(doc, "Víctor", "Análisis crítico desde los derechos digitales y la proporcionalidad penal")

    # --- Introducción ---
    add_heading_styled(doc, "Introducción", level=1)
    doc.add_paragraph(
        'El caso presentado en el artículo de El Espectador pone de manifiesto una de las tensiones '
        'más relevantes de la era digital: el conflicto entre la protección de los derechos de autor '
        'y el derecho al acceso al conocimiento. Diego Gómez, un joven biólogo colombiano, compartió '
        'una tesis académica en la plataforma Scribd con el propósito de facilitar la investigación '
        'científica entre sus compañeros. Esta acción, motivada por la buena fe y la carencia de '
        'recursos para acceder a bases de datos especializadas, lo llevó a enfrentar un proceso penal '
        'que podría haberlo condenado hasta por ocho años de prisión.'
    )

    # --- Pregunta 1 ---
    add_heading_styled(doc, "1. Identificación de los posibles delitos", level=1)

    add_heading_styled(doc, "a) Violación a los derechos patrimoniales de autor (delito principal):", level=2)
    doc.add_paragraph(
        'Al subir la tesis a Scribd sin contar con la autorización del autor, Gómez habría incurrido '
        'en la vulneración de los derechos exclusivos del titular sobre la obra. En este caso:'
    )
    add_bullet(doc, "Se realizó una copia digital del documento original.", "Reproducción: ")
    add_bullet(doc, "La tesis fue puesta a disposición del público en una plataforma de alcance global.", "Distribución: ")
    add_bullet(doc, "Al cargar el archivo en Scribd, se habilitó el acceso para cualquier persona con conexión a internet.", "Comunicación pública: ")

    add_heading_styled(doc, "b) Aspectos que atenúan la conducta:", level=2)
    add_bullet(doc, "No existió ánimo de lucro: en ningún momento obtuvo beneficio económico.")
    add_bullet(doc, "Finalidad académica legítima: el propósito era facilitar la investigación científica.")
    add_bullet(doc, "La tesis ya circulaba en internet: el documento estaba previamente disponible.")
    add_bullet(doc, "No hubo apropiación de autoría: Gómez nunca se atribuyó la autoría de la obra.")

    doc.add_paragraph(
        'Desde una perspectiva de derechos digitales, la criminalización de esta conducta resulta '
        'cuestionable, pues equipara a un estudiante que comparte conocimiento con un pirata comercial '
        'que lucra con obras ajenas.'
    )

    # --- Pregunta 2 ---
    add_heading_styled(doc, "2. Leyes y artículos presuntamente infringidos", level=1)

    add_heading_styled(doc, "Legislación Nacional", level=2)

    add_bold_paragraph(doc, "Artículo 271 del Código Penal (modificado por Ley 1032 de 2006):")
    doc.add_paragraph(
        'Sanciona a quien "por cualquier medio o procedimiento, sin autorización previa y expresa '
        'del titular, reproduzca, traduzca, distribuya, comunique, o por cualquier otro medio, ponga '
        'a disposición del público una obra protegida". La reforma de 2006 fue especialmente problemática '
        'porque eliminó el requisito de ánimo de lucro, ampliando excesivamente las conductas penalizables.'
    )

    add_bold_paragraph(doc, "Ley 23 de 1982 (Ley de Derechos de Autor):")
    doc.add_paragraph(
        'Establece los derechos morales y patrimoniales que protegen a los creadores. Aunque contempla '
        'algunas excepciones (derecho de cita, uso personal), son insuficientes para el uso académico '
        'colaborativo en entornos digitales.'
    )

    add_bold_paragraph(doc, "Ley 44 de 1993:")
    doc.add_paragraph('Complementa la Ley 23 de 1982 y refuerza la protección de derechos patrimoniales.')

    add_heading_styled(doc, "Legislación Internacional Aplicable", level=2)
    add_bullet(doc, "Régimen Común sobre Derecho de Autor para la Comunidad Andina.", "Decisión Andina 351 de 1993: ")
    add_bullet(doc, "Establece estándares mínimos de protección a los que Colombia está adherida.", "Convenio de Berna: ")

    add_heading_styled(doc, "Análisis crítico de la normativa", level=2)
    doc.add_paragraph(
        'La principal falencia de este marco legal radica en que no distingue entre piratería comercial '
        'y compartir académico. La Ley 1032 de 2006, al eliminar el requisito de ánimo de lucro del '
        'artículo 271, convirtió en delito penal cualquier forma de reproducción no autorizada, sin '
        'importar el contexto, la finalidad ni el impacto económico real.'
    )

    # --- Pregunta 3 ---
    add_heading_styled(doc, "3. Penas aplicables en caso de condena", level=1)

    add_heading_styled(doc, "Penas privativas de la libertad", level=2)
    doc.add_paragraph(
        'Prisión de 4 a 8 años. Para poner esto en perspectiva:'
    )
    add_bullet(doc, "El hurto simple (Art. 239) tiene pena de 2 a 6 años.")
    add_bullet(doc, "Las lesiones personales (Art. 111) pueden tener penas inferiores.")
    doc.add_paragraph(
        'Esto significa que compartir una tesis podría haberse castigado con mayor severidad '
        'que robar o causar daño físico a otra persona.'
    )

    add_heading_styled(doc, "Penas pecuniarias", level=2)
    doc.add_paragraph(
        'Multa de 26.66 a 1.000 SMLMV, lo que en valores actuales (2026) representaría entre '
        'aproximadamente $37.9 millones y $1.423 millones de pesos colombianos, una cifra devastadora '
        'para un estudiante universitario.'
    )

    add_heading_styled(doc, "Penas accesorias", level=2)
    add_bullet(doc, "Antecedentes penales permanentes.")
    add_bullet(doc, "Posible reparación civil a favor del titular.")
    add_bullet(doc, "Inhabilitación para funciones públicas.")

    add_heading_styled(doc, "Reflexión sobre la proporcionalidad", level=2)
    doc.add_paragraph(
        'El movimiento #CompartirNoEsDelito, liderado por la Fundación Karisma y apoyado por '
        'organizaciones internacionales como Creative Commons y la EFF, denunció la desproporción '
        'entre la conducta y la sanción. Esta campaña fue decisiva para generar un debate público '
        'que eventualmente condujo a reformas legislativas.'
    )

    # --- Pregunta 4 ---
    add_heading_styled(doc, "4. Recomendaciones como experto en seguridad informática", level=1)

    add_heading_styled(doc, "A. Para la comunidad académica y los ciudadanos", level=2)
    add_bullet(doc, "Las universidades deben contar con protocolos claros sobre cómo compartir y distribuir "
               "material académico, con oficinas de propiedad intelectual que asesoren gratuitamente.",
               "Implementar políticas institucionales de propiedad intelectual: ")
    add_bullet(doc, "Promover que los autores publiquen bajo licencias Creative Commons u otras licencias "
               "abiertas que permitan la redistribución con atribución.",
               "Fomentar el uso de licencias abiertas: ")
    add_bullet(doc, "Emplear repositorios institucionales, plataformas de acceso abierto y bibliotecas "
               "digitales autorizadas.",
               "Utilizar canales legales de difusión: ")
    add_bullet(doc, "Incluir en la educación superior módulos obligatorios sobre derechos de autor y los "
               "riesgos legales del uso de material protegido en el entorno digital.",
               "Formación en alfabetización digital y jurídica: ")

    add_heading_styled(doc, "B. Para el Estado colombiano", level=2)
    add_bullet(doc, "Reintroducir el criterio de ánimo de lucro como elemento del tipo penal, o crear "
               "excepciones explícitas para el uso académico.",
               "Distinguir entre piratería comercial y uso académico: ")
    add_bullet(doc, "Las disputas de propiedad intelectual académica deberían resolverse mediante "
               "mecanismos civiles o administrativos.",
               "Priorizar la vía civil sobre la penal: ")
    add_bullet(doc, "Implementar legislación que obligue a que toda investigación financiada con recursos "
               "públicos sea de acceso abierto.",
               "Fortalecer el acceso abierto: ")
    add_bullet(doc, "Si bien la Ley 1915 de 2018 introdujo avances, la legislación sigue requiriendo una "
               "reforma integral para el entorno digital.",
               "Modernizar la Ley 23 de 1982: ")
    add_bullet(doc, "Crear un organismo que monitoree conflictos entre derechos de autor y acceso al "
               "conocimiento, proponiendo reformas basadas en evidencia.",
               "Establecer un observatorio de propiedad intelectual digital: ")

    # --- Conclusión ---
    doc.add_page_break()
    add_heading_styled(doc, "Conclusión", level=1)
    doc.add_paragraph(
        'El caso de Diego Gómez demuestra que la legislación colombiana sobre derechos de autor fue '
        'diseñada para un mundo analógico y ha sido insuficientemente actualizada para las dinámicas '
        'del ecosistema digital. Criminalizar a un estudiante que comparte conocimiento sin ánimo de '
        'lucro con las mismas herramientas legales contra la piratería comercial no solo es '
        'desproporcionado, sino que desincentiva la investigación y el acceso al conocimiento.'
    )

    # --- Referencias ---
    add_heading_styled(doc, "Referencias", level=1)
    refs = [
        'El Espectador (2014). "Estudiante podría ir a cárcel por divulgar tesis en una plataforma".',
        "Código Penal Colombiano – Ley 599 de 2000, Artículo 271.",
        "Ley 23 de 1982 – Sobre Derechos de Autor.",
        "Ley 1032 de 2006 – Modificación del Artículo 271 del Código Penal.",
        "Ley 1915 de 2018 – Modernización de Derechos de Autor.",
        "Decisión Andina 351 de 1993.",
        "Fundación Karisma – Caso Diego Gómez: #CompartirNoEsDelito.",
        "Creative Commons Colombia.",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    path = os.path.join(ENTREGA_DIR, "Solucion_Eje4_Victor.docx")
    doc.save(path)
    print(f"✅ Generado: {path}")


# ==============================================================================
# ARCHIVO 3 - HAROLD (Gobernanza y prevención institucional)
# ==============================================================================
def generate_harold():
    doc = Document()
    set_normal_style(doc)
    add_title_page(doc, "Harold", "Gobernanza de seguridad de la información y prevención institucional")

    # --- Presentación ---
    add_heading_styled(doc, "Presentación del Caso", level=1)
    doc.add_paragraph(
        'El artículo periodístico de El Espectador da cuenta de la situación de Diego Gómez, un '
        'estudiante de biología de la Universidad del Quindío, quien enfrentó un proceso penal por '
        'haber compartido en la plataforma Scribd una tesis de maestría sin la debida autorización '
        'del autor original. Este caso resulta especialmente relevante desde la óptica de la seguridad '
        'de la información, ya que evidencia cómo el desconocimiento de los marcos legales puede '
        'exponer a estudiantes, profesionales e instituciones a riesgos jurídicos significativos.'
    )

    # --- Pregunta 1 ---
    add_heading_styled(doc, "1. Delitos en los que posiblemente incurrió el ciudadano", level=1)

    add_heading_styled(doc, "1.1. Violación a los derechos patrimoniales de autor", level=2)
    doc.add_paragraph(
        'Este es el tipo penal central del caso. Los derechos patrimoniales confieren al autor la '
        'exclusividad para explotar económicamente su obra. Al subir la tesis a Scribd, Diego Gómez '
        'realizó las siguientes acciones sin autorización:'
    )
    add_bullet(doc, "Creó una copia del documento en formato digital.", "Reproducción digital: ")
    add_bullet(doc, "Lo hizo accesible a millones de usuarios a nivel mundial.", "Puesta a disposición del público: ")
    add_bullet(doc, "Compartió el enlace en un grupo de Facebook, ampliando su alcance.", "Distribución no autorizada: ")

    add_heading_styled(doc, "1.2. Consideraciones sobre la ausencia de otros delitos", level=2)
    doc.add_paragraph('Es importante aclarar que en este caso no se configuran otros delitos frecuentemente '
                      'asociados al entorno digital:')
    add_bullet(doc, "Gómez no se atribuyó la autoría de la tesis.", "No hay plagio: ")
    add_bullet(doc, "No ingresó ilícitamente a ningún sistema (Art. 269A, Ley 1273 de 2009).",
               "No hay acceso abusivo: ")
    add_bullet(doc, "No empleó herramientas tecnológicas ilícitas.", "No hay uso de software malicioso: ")
    add_bullet(doc, "El documento fue obtenido de fuentes públicas en internet.",
               "No hay interceptación de comunicaciones: ")

    doc.add_paragraph(
        'Desde la perspectiva de la gobernanza de seguridad de la información, este caso demuestra '
        'que los riesgos legales no siempre provienen de acciones técnicas sofisticadas, sino que '
        'pueden originarse en actividades cotidianas como compartir un archivo cuando no se cuenta '
        'con el conocimiento adecuado sobre las regulaciones vigentes.'
    )

    # --- Pregunta 2 ---
    add_heading_styled(doc, "2. Marco legal infringido: Leyes y artículos aplicables", level=1)

    add_heading_styled(doc, "2.1. Normativa penal", level=2)
    add_table(doc,
        ["Norma", "Disposición", "Contenido relevante"],
        [
            ["Código Penal\n(Ley 599/2000)", "Art. 271",
             "Tipifica la violación a los derechos patrimoniales de autor. Pena: 4-8 años de prisión + multa."],
            ["Ley 1032 de 2006", "Modifica Art. 271",
             "Endureció las penas y eliminó el ánimo de lucro como requisito del tipo penal."],
        ]
    )

    add_heading_styled(doc, "2.2. Normativa especial de derechos de autor", level=2)
    add_table(doc,
        ["Norma", "Disposición", "Contenido relevante"],
        [
            ["Ley 23 de 1982", "Arts. 2, 3, 12, 72",
             "Ley marco de derechos de autor en Colombia. Protege obras científicas, literarias y artísticas."],
            ["Ley 44 de 1993", "Arts. 51-52",
             "Complementa la Ley 23/1982 respecto a derechos patrimoniales y morales."],
            ["Ley 1915 de 2018", "Múltiples arts.",
             "Moderniza el régimen. Introduce excepciones para bibliotecas y archivos (posterior al caso)."],
        ]
    )

    add_heading_styled(doc, "2.3. Normativa internacional", level=2)
    add_table(doc,
        ["Norma", "Ámbito", "Relevancia"],
        [
            ["Decisión Andina 351/1993", "Comunidad Andina",
             "Régimen supranacional de derechos de autor, aplicable directamente en Colombia."],
            ["Convenio de Berna (1886)", "Internacional",
             "Establece estándares mínimos de protección."],
            ["Tratado OMPI (WCT, 1996)", "Internacional",
             "Aborda la protección de obras en el entorno digital, incluyendo puesta a disposición del público."],
        ]
    )

    add_heading_styled(doc, "2.4. Análisis desde la gestión de riesgos", level=2)
    doc.add_paragraph(
        'Desde la perspectiva de un Sistema de Gestión de Seguridad de la Información (SGSI), este '
        'caso evidencia la necesidad de incluir el cumplimiento legal y regulatorio como pilar '
        'fundamental. La norma ISO/IEC 27001 (Anexo A, control A.18) exige que las organizaciones '
        'identifiquen y cumplan los requisitos legales aplicables, incluyendo la propiedad intelectual.'
    )

    # --- Pregunta 3 ---
    add_heading_styled(doc, "3. Penas aplicables en caso de condena", level=1)

    add_heading_styled(doc, "3.1. Penas principales", level=2)
    add_table(doc,
        ["Tipo de pena", "Detalle", "Observación"],
        [
            ["Prisión", "4 a 8 años", "Pena intramural en establecimiento penitenciario"],
            ["Multa", "26.66 a 1.000 SMLMV", "Equivalente a aprox. $37.9M – $1.423M COP (2026)"],
        ]
    )

    add_heading_styled(doc, "3.2. Penas accesorias y consecuencias derivadas", level=2)
    add_bullet(doc, "Registro permanente que dificulta el acceso a empleos, visas y contratos.",
               "Antecedentes judiciales: ")
    add_bullet(doc, "Por tiempo equivalente a la pena de prisión.",
               "Inhabilitación para funciones públicas: ")
    add_bullet(doc, "Obligación de reparar económicamente al titular de los derechos vulnerados.",
               "Indemnización de perjuicios: ")
    add_bullet(doc, "Los equipos informáticos utilizados podrían haber sido decomisados.",
               "Decomiso de bienes: ")

    add_heading_styled(doc, "3.3. Impacto desde la perspectiva de seguridad de la información", level=2)
    doc.add_paragraph(
        'Este caso ilustra cómo un incidente de seguridad de la información no siempre involucra un '
        'ataque técnico. La filtración no autorizada de información protegida puede tener consecuencias '
        'legales tan severas como un ciberataque. En la gestión de la seguridad de la información, '
        'la clasificación de activos de información y el control de su distribución son fundamentales. '
        'Una tesis de grado es un activo de información que tiene un propietario y cuya distribución '
        'está regulada.'
    )

    # --- Pregunta 4 ---
    add_heading_styled(doc, "4. Recomendaciones como experto en seguridad de la información", level=1)

    add_heading_styled(doc, "A. Recomendaciones para las instituciones educativas (enfoque preventivo)", level=2)
    add_bullet(doc, "Basado en la norma ISO/IEC 27001, incluyendo políticas de protección de propiedad "
               "intelectual y clasificación de activos de información.",
               "1. Implementar un SGSI: ")
    add_bullet(doc, "Que establezca claramente qué pueden y qué no pueden hacer los miembros de la "
               "comunidad universitaria con los recursos académicos digitales.",
               "2. Desarrollar una Política de Uso Aceptable (AUP): ")
    add_bullet(doc, "DRM para documentos sensibles, marcas de agua digitales para rastreo, "
               "y gestión de acceso basada en roles (RBAC).",
               "3. Implementar controles técnicos en repositorios: ")
    add_bullet(doc, "Programas anuales de formación sobre propiedad intelectual y seguridad de la "
               "información con casos prácticos como el de Diego Gómez.",
               "4. Capacitación obligatoria: ")
    add_bullet(doc, "Un servicio institucional gratuito donde estudiantes e investigadores consulten "
               "antes de publicar o compartir material de terceros.",
               "5. Oficina de asesoría en propiedad intelectual: ")

    add_heading_styled(doc, "B. Recomendaciones para el Estado (enfoque de política pública)", level=2)
    add_bullet(doc, "Crear una escala de sanciones que considere intencionalidad, ánimo de lucro, "
               "impacto económico real y contexto. Las sanciones penales deben reservarse para "
               "piratería organizada.",
               "1. Establecer un marco de respuesta proporcional: ")
    add_bullet(doc, "Legislar para que la investigación financiada con recursos públicos sea obligatoriamente "
               "de acceso abierto.",
               "2. Desarrollar un marco legal para el acceso abierto: ")
    add_bullet(doc, "Un organismo administrativo especializado en disputas de propiedad intelectual "
               "en el entorno digital, más ágil y proporcional que el sistema penal.",
               "3. Crear una entidad especializada de resolución de conflictos: ")
    add_bullet(doc, "Alinear la normativa con las mejores prácticas internacionales como el 'fair use' "
               "estadounidense o las excepciones europeas para investigación.",
               "4. Armonizar la legislación con estándares internacionales: ")
    add_bullet(doc, "Programas nacionales de alfabetización jurídico-digital que permitan a los ciudadanos "
               "conocer sus derechos y obligaciones en el entorno digital.",
               "5. Promover la ciberseguridad jurídica: ")

    # --- Conclusión ---
    doc.add_page_break()
    add_heading_styled(doc, "Conclusión", level=1)
    doc.add_paragraph(
        'El caso Diego Gómez constituye un punto de inflexión en la discusión sobre derechos de autor '
        'y acceso al conocimiento en Colombia. Desde la gobernanza de la seguridad de la información, '
        'este caso nos enseña que la protección de los activos de información no se limita a aspectos '
        'técnicos (firewalls, cifrado, antivirus), sino que abarca dimensiones legales, éticas y de '
        'política pública que deben ser gestionadas de manera integral. Las organizaciones e instituciones '
        'educativas deben adoptar un enfoque proactivo que combine controles técnicos, políticas '
        'institucionales y formación continua para prevenir situaciones similares.'
    )

    # --- Referencias ---
    add_heading_styled(doc, "Referencias", level=1)
    refs = [
        'El Espectador (2014). "Estudiante podría ir a cárcel por divulgar tesis en una plataforma".',
        "Código Penal Colombiano – Ley 599 de 2000, Art. 271.",
        "Ley 23 de 1982 – Sobre Derechos de Autor (Colombia).",
        "Ley 1032 de 2006 – Modificación del Art. 271 del Código Penal.",
        "Ley 44 de 1993 – Complementaria de Derechos de Autor.",
        "Ley 1915 de 2018 – Modernización de Derechos de Autor.",
        "Ley 1273 de 2009 – Delitos Informáticos en Colombia.",
        "Decisión Andina 351 de 1993 – Régimen Común sobre Derecho de Autor.",
        "ISO/IEC 27001:2022 – Sistema de Gestión de Seguridad de la Información.",
        "Fundación Karisma – Campaña #CompartirNoEsDelito.",
        "Tratado OMPI sobre Derecho de Autor (WCT, 1996).",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    path = os.path.join(ENTREGA_DIR, "Solucion_Eje4_Harold.docx")
    doc.save(path)
    print(f"✅ Generado: {path}")


# ==============================================================================
# MAIN
# ==============================================================================
if __name__ == "__main__":
    print("Generando archivos .docx para Actividad Evaluativa Eje 4...\n")
    generate_melqui()
    generate_victor()
    generate_harold()
    print("\n🎉 ¡Todos los archivos han sido generados exitosamente!")
