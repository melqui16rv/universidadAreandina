from docx import Document
import os

doc = Document()
doc.add_heading('Actividad Evaluativa Eje 2 - Victor', level=1)

doc.add_paragraph('Desarrollo del cuestionario basado en el documental "Delitos Informáticos" y el material de estudio referente a la superación de las limitaciones autoimpuestas en nuestra seguridad y el impacto de la Ingeniería Social.')

preguntas = [
    "Identificar la modalidad de delitos informáticos que se mencionan en el video.",
    "Tipo de población afectada y porque se incurre en estos eventos.",
    "¿Cuál es la entidad que recopila las evidencias y que técnicas utiliza?",
    "¿Bajo qué ley se castiga al delincuente informático?",
    "Usted como experto en seguridad que recomendaciones haría para que no sucedan estos casos.",
    "¿Qué acciones debe tomar el estado al respecto?",
    "¿En qué consiste el enmascarado de una IP?"
]

respuestas_victor = [
    "Se mencionan modalidades que van desde el nivel de usuario (fraude, chantaje, robos de información y suplantación de identidad) hasta ataques orquestados a gran escala como infiltración en sistemas del Estado (como Anonymous), generación de pánico económico contra bancos y mercados clandestinos en la Deep Web para comercializar material ilegal.",
    "Se afecta principalmente a personas que actúan con exceso de confianza o descuido digital. Incurren en estos hechos porque la ingeniería social es muy efectiva; el atacante explota nuestros miedos y auto-limitaciones para engañarnos haciéndonos abrir correos con malware o revelar nuestras credenciales a perfiles que ni siquiera conocemos.",
    "El ente que interviene inicialmente tras una denuncia es la Fiscalía General de la Nación. Se apoyan en la disciplina de la informática forense, utilizando hardware especializado como el escáner o duplicador forense que incluye una protección contra escritura (write-blocker). Esto asegura que se capture una copia forense sin comprometer el disco o elemento original.",
    "Históricamente se catalogaban estos casos como fraudes o violaciones de derechos de autor, pero actualmente existen leyes específicas para sancionarlo (como la Ley 1273). Las medidas varían desde condenas de 0 a 4 años hasta sentencias más severas de 8 a 12 años de prisión, agravándose por el daño causado al ciudadano o si el perpetrador actuó en asociación criminal.",
    "A nivel corporativo, es vital aplicar Ethical Hacking para diagnosticar fallas mediante intrusiones controladas. A nivel personal, el usuario debe erradicar la postura de que 'la seguridad no es asunto suyo', asumiendo su responsabilidad al no abrir correos de remitentes dudosos y cuidando fervientemente el rastro fotográfico o confidencial que deposita en sus redes sociales.",
    "El Estado está obligado a tomar en serio las denuncias para formalizar a los afectados como verdaderas víctimas y dar paso a indagaciones contundentes. Aunado a esto, las autoridades gubernamentales deben legislar y renovar las leyes penales constantemente, pues el cibercrimen siempre está creando nuevas vías para evadir la justicia.",
    "El enmascaramiento de IP es una táctica disuasiva ejecutada regularmente a través de plataformas como Tor. Su finalidad principal es rebotar la señal de internet para que un ciberdelincuente evite el rastreo geográfico por completo; permitiéndole, por ejemplo, operar su computadora en Colombia mientras figura virtualmente como si estuviese conectado en un país europeo."
]

for p, r in zip(preguntas, respuestas_victor):
    p_bold = doc.add_paragraph()
    p_bold.add_run(p).bold = True
    doc.add_paragraph(r)

file_path = os.path.join(r"c:\github\universidadAreandina\semestre7\corte2\ASPECTOS LEGALES DE LA SEGURID\eje2", "activida", "Actividad_Evaluativa_Eje2_Victor.docx")
doc.save(file_path)
print(f"Documento de Victor guardado en {file_path}")