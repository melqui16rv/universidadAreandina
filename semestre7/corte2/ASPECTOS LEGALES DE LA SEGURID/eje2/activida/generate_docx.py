from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Titulo Nivel 1
heading = doc.add_heading('Solución Actividad - Eje 2: Delitos Informáticos', level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduccion
doc.add_paragraph('A continuación se presenta el análisis y desarrollo de la actividad sobre Delitos Informáticos, '
                  'basado en el documental suministrado.', style='BodyText')

doc.add_heading('Análisis de Preguntas', level=2)

# Pregunta 1
p1 = doc.add_paragraph()
p1.add_run('1. Modalidad de delitos informáticos que se mencionan en el video:').bold = True
p = doc.add_paragraph('Dentro del video se mencionan varios tipos de delitos, desde los más tradicionales hasta los ejecutados en la "Deep Web" o internet profunda:')
p.add_run('\n- Fraude, suplantación de identidad, chantaje y robo de información.').italic = True
p.add_run('\n- Interceptaciones ilícitas y asaltos informáticos a sistemas gubernamentales/militares.').italic = True
p.add_run('\n- Pánico económico (por ejemplo, el caso contra una entidad bancaria para descapitalizarla).').italic = True
p.add_run('\n- Tráfico ilegal en la Deep Web: venta de drogas, armas, visas falsas, pasaportes, órganos y material con contenido de abuso infantil.').italic = True

# Pregunta 2
p2 = doc.add_paragraph()
p2.add_run('2. Tipo de población afectada y por qué se incurre en estos eventos:').bold = True
doc.add_paragraph('La población más afectada son aquellas personas que utilizan de forma inocente o ingenua las tecnologías de la información. Suelen ser individuos que carecen de cultura informática, abren correos sin verificar el remitente o aceptan invitaciones en redes sociales de desconocidos. Incurren en estos eventos debido al desconocimiento de los riesgos (Ingeniería Social) y al exceso de confianza.')

# Pregunta 3
p3 = doc.add_paragraph()
p3.add_run('3. ¿Cuál es la entidad que recopila las evidencias y qué técnicas utiliza?').bold = True
doc.add_paragraph('Al realizarse una denuncia, el Estado asigna la labor investigativa a la Fiscalía General de la Nación. '
                  'Las técnicas utilizadas corresponden a la "Informática Forense". Para no alterar las evidencias se emplean, '
                  'entre otras cosas, aparatos físicos como el duplicador forense con bloqueo contra escritura (write-blocker). '
                  'Esto permite obtener y evaluar una imagen forense del dispositivo vulnerado sin modificar o dañar los datos originales del disco del afectado.')

# Pregunta 4
p4 = doc.add_paragraph()
p4.add_run('4. ¿Bajo qué ley se castiga al delincuente informático?').bold = True
doc.add_paragraph('El video hace mención (aparentemente referenciando la Ley 1273 en Colombia) a que anteriormente los delitos informáticos se penalizaban tipificando las conductas bajo leyes de derecho de autor, falsedad o suplantación. Las condenas impuestas actualmente dependen del grado y la modalidad, contemplando sentencias de entre 4 a 8 años o de 8 a 12 años, entre otras agravantes y dependiendo de si pertenecen a una organización criminal.')

# Pregunta 5
p5 = doc.add_paragraph()
p5.add_run('5. Usted como experto en seguridad, ¿qué recomendaciones haría para que no sucedan estos casos?').bold = True
doc.add_paragraph('Basado en la información recopilada, se sugieren las siguientes recomendaciones:')
p_rec = doc.add_paragraph(style='List Bullet')
p_rec.add_run('A nivel Corporativo e Institucional: ').bold = True
p_rec.add_run('Realizar pruebas de auditoría de seguridad preventiva (Ethical Hacking) para identificar vulnerabilidades y parchar agujeros de seguridad antes de que sean explotados por cibercriminales.')
p_rec2 = doc.add_paragraph(style='List Bullet')
p_rec2.add_run('A nivel Personal/Usuarios: ').bold = True
p_rec2.add_run('Controlar estrictamente la información que se sube en redes sociales asumiendo que es de dominio público; verificar los remitentes de los correos electrónicos; no abrir archivos de dudosa procedencia (malware/phishing) y no entablar conversaciones ni aceptar conexiones de gente desconocida.')

# Pregunta 6
p6 = doc.add_paragraph()
p6.add_run('6. ¿Qué acciones debe tomar el Estado al respecto?').bold = True
doc.add_paragraph('El Estado debe recepcionar eficazmente las denuncias para que se inicie un proceso de investigación reconociendo debidamente al afectado como víctima. Asimismo, el marco legal debe estar en constante actualización, aprobando normativas que penalicen las nuevas formas de modalidades delictivas a medida que la tecnología evoluciona, ya que los atacantes buscan maneras cada vez más sofisticadas de borrar sus rastros.')

# Pregunta 7
p7 = doc.add_paragraph()
p7.add_run('7. ¿En qué consiste el enmascarado de una IP?').bold = True
doc.add_paragraph('Consiste en el uso de herramientas o sistemas de enrutamiento (como Tor y sus múltiples capas/nodos) que encubren la dirección IP de origen real de la computadora del usuario. Esto permite a un cibercriminal ocultar el rastro de su ubicación física; por ejemplo, el delincuente puede encontrarse en Colombia desarrollando sus ataques, pero el sistema registrará una procedencia aparente en Francia o Alemania, evadiendo así los sistemas de rastreo tradicionales.')

# Guardar documento
file_path = os.path.join(r"c:\github\universidadAreandina\semestre7\corte2\ASPECTOS LEGALES DE LA SEGURID\eje2", "activida", "Solucion_Actividad_Eje2.docx")
doc.save(file_path)
print(f"Documento guardado en {file_path}")
