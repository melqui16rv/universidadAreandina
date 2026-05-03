from docx import Document
import os

doc = Document()
doc.add_heading('Solución Actividad Eje 2 - Harold', level=1)

doc.add_paragraph('En desarrollo de la actividad se responden las preguntas planteadas en la actividad, tomando como base el video documental y el referente de pensamiento del módulo.')

preguntas = [
    "Identificar la modalidad de delitos informáticos que se mencionan en el video.",
    "Tipo de población afectada y porque se incurre en estos eventos.",
    "¿Cuál es la entidad que recopila las evidencias y que técnicas utiliza?",
    "¿Bajo qué ley se castiga al delincuente informático?",
    "Usted como experto en seguridad que recomendaciones haría para que no sucedan estos casos.",
    "¿Qué acciones debe tomar el estado al respecto?",
    "¿En qué consiste el enmascarado de una IP?"
]

respuestas_harold = [
    "En el documental se exponen ataques iniciales o tradicionales como la suplantación de identidad, el fraude, el robo de información y el chantaje. Además, se detallan crímenes de mayor magnitud como las interceptaciones a sistemas estatales, pánico económico bancario y delitos originados en la Deep Web (tráfico de armas, drogas y pornografía infantil). Estos últimos demuestran el alto grado de especialización que alcanzan algunas bandas criminales en la actualidad.",
    "La población afectada son ciudadanos comunes que usan la tecnología de forma inocente o con falta de cultura informática. Caen en estos eventos producto de la ingeniería social, abriendo correos dudosos o aceptando solicitudes extrañas. Como señala nuestro referente, son los temores y limitaciones infundadas en torno a la tecnología lo que nos hace vulnerables. Generalmente, el exceso de confianza y el desconocimiento facilitan enormemente la labor de los estafadores digitales.",
    "La entidad encargada de liderar la investigación es la Fiscalía General de la Nación. Aplican técnicas de informática forense mediante dispositivos como el duplicador forense. Esta herramienta emplea un bloqueo de escritura (write blocker) para tomar una imagen forense exacta del disco sin alterar la prueba original. Esto garantiza la total integridad de la cadena de custodia y permite un análisis detallado en laboratorios certificados.",
    "Previamente se recurría a leyes de derechos de autor o falsedad, pero la ley aplicable actualmente (se sugiere la Ley 1273) establece penas específicas. Las sentencias varían dependiendo de la gravedad (0 a 4 años, 4 a 8 años, o de 8 a 12 años) y si el perpetrador forma parte de una red criminal. Así, el sistema judicial busca castigar con mayor rigor el impacto generado en la sociedad civil.",
    "Recomiendo a las empresas contratar servicios de Ethical Hacking para blindar sus sistemas preventivamente. A los ciudadanos, les aconsejo controlar lo que publican en redes sociales, verificar siempre la identidad de los remitentes y no subestimar la ingeniería social. Hay que distinguir entre limitaciones reales y temores, y educarse responsablemente. Solo a través de la concienciación y el autocuidado lograremos mitigar estos riesgos de manera efectiva y sostenible.",
    "El Estado tiene el deber de asegurar que los denunciantes reciban estatus de víctimas, activando inmediatamente protocolos de investigación. De la misma manera, debe actualizar el Código Penal y normatividades para ir a la par de las nuevas metodologías cibernéticas que borran rastros cada vez con más eficacia. Una rápida judicialización es clave para disuadir a futuras organizaciones dedicadas al cibercrimen y proteger al ciudadano.",
    "Consiste en el uso de redes de enrutadores (como la herramienta Tor) para que la conexión no pueda ser rastreada. Mediante esta técnica, la dirección IP de origen de un atacante en Colombia puede ser enmascarada para que el sistema la identifique en Francia o Alemania, protegiendo su anonimato. Todo esto complica significativamente la labor de rastreo, exigiendo peritos informáticos mucho más capacitados e interconectados."
]

for p, r in zip(preguntas, respuestas_harold):
    p_bold = doc.add_paragraph()
    p_bold.add_run(p).bold = True
    doc.add_paragraph(r)

file_path = os.path.join(r"c:\github\universidadAreandina\semestre7\corte2\ASPECTOS LEGALES DE LA SEGURID\eje2", "activida", "Solucion_Actividad_Eje2_Harold.docx")
doc.save(file_path)
print(f"Documento de Harold actualizado guardado en {file_path}")